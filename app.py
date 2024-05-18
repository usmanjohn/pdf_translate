import openai
import streamlit as st
from docx import Document
from io import BytesIO
from pdf2image import convert_from_path
from pytesseract import image_to_string
from autocorrect import Speller
import os

# Set up OpenAI API key
openai.api_key = st.secrets['my_key']

# Function to translate text
def translate_text(text, target_language="ko"):
    prompt = f"Translate the following text to {target_language}:\n\n{text}"
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a professional insurance lawyer who is good at translating documents accurately while keeping paragraphing and styles."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000
    )
    translated_text = response.choices[0].message.content.strip()
    return translated_text

# Function to translate a document
def translate_document(doc, target_language="ko"):
    translated_doc = Document()
    for paragraph in doc.paragraphs:
        translated_text = translate_text(paragraph.text, target_language)
        translated_doc.add_paragraph(translated_text)
    return translated_doc

# OCR function to convert PDF to text
def ocr_pdf_to_text(pdf_path):
    images = convert_from_path(pdf_path)
    text = ""
    for image in images:
        text += image_to_string(image)
    return text

# Correct text using autocorrect
def correct_text(text):
    spell = Speller()
    corrected_text = spell(text)
    return corrected_text


import re

def preprocess_text(text):
    # Remove control characters and NULL bytes
    text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)
    return text
# Streamlit interface
st.title("Document Translator")

uploaded_file = st.file_uploader("Upload a scanned PDF document", type="pdf")

if uploaded_file is not None:
    with open("temp.pdf", "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    st.write("Converting PDF to text...")
    extracted_text = ocr_pdf_to_text("temp.pdf")

    st.write("Cleaning text...")
    cleaned_text = preprocess_text("temp.pdf")

    
    st.write("Correcting text errors...")
    corrected_text = correct_text(cleaned_text)
    
    st.write("Refining text with OpenAI...")
   
    
    doc = Document()
    doc.add_paragraph(corrected_text)
    
    target_language = st.selectbox("Select target language", [("Korean", "ko"), ("English", "eng")], format_func=lambda x: x[0])
    
    if st.button("Translate"):
        translated_doc = translate_document(doc, target_language)
        output = BytesIO()
        translated_doc.save(output)
        output.seek(0)
        
        st.download_button(
            label="Download Translated Document",
            data=output,
            file_name="translated_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    os.remove("temp.pdf")
